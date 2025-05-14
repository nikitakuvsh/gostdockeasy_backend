import logging
import os
import random
import string
from fastapi import FastAPI, UploadFile, Form, Depends, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy.ext.asyncio import create_async_engine, async_sessionmaker, AsyncSession
from sqlalchemy.orm import DeclarativeBase
from sqlalchemy.future import select
from fastapi.responses import FileResponse
from docx import Document
from docx2pdf import convert
from models import Faculty
from fastapi import APIRouter
from sqlalchemy import extract, func
from datetime import datetime
from database import get_session
from models import Submission
import docx.shared
import shutil
import uvicorn

# Настроим логирование
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

DATABASE_URL = "postgresql+asyncpg://gostdockeasy_user:tri4Wq1P0hKuUTAOi0jMx41nHoyGYftL@dpg-d0ido1mmcj7s73dif520-a:5432/gostdockeasy"

# Создание асинхронного движка базы данных
engine = create_async_engine(DATABASE_URL, echo=True)
AsyncSessionLocal = async_sessionmaker(engine, expire_on_commit=False)

# Абсолютный путь для временной директории
TEMP_DIR = os.path.join(os.getcwd(), "temp_files")
os.makedirs(TEMP_DIR, exist_ok=True)

# Базовый класс для всех моделей
class Base(DeclarativeBase):
    pass

# Настройки CORS для разрешения всех источников
app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Получение сессии
async def get_session() -> AsyncSession:
    async with AsyncSessionLocal() as session:
        yield session

# Создание таблиц при запуске
@app.on_event("startup")
async def on_startup():
    async with engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)

# Функция для добавления нижнего колонтитула
def add_footer(doc: Document, faculty: str):
    # Добавляем нижний колонтитул
    section = doc.sections[-1]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    
    # Добавляем текст в нижний колонтитул
    run = paragraph.add_run(f"Тестовый нижний колонтитул: {faculty}")
    run.font.size = 140000  # размер шрифта
    
    # Можно настроить расположение, например, выровнять по центру:
    paragraph.alignment = 1  # 0 - влево, 1 - по центру, 2 - вправо

# Функция для генерации уникального имени файла
def generate_unique_filename(extension=".docx"):
    return ''.join(random.choices(string.ascii_lowercase + string.digits, k=10)) + extension

# Функция для заполнения шаблона Word
def fill_template(file_path: str, faculty: str):
    logger.info(f"Файл шаблона: {file_path}")
    doc = Document(file_path)

    # Установка параметров страницы
    section = doc.sections[0]
    section.top_margin = 2_000_000     # 2 см
    section.bottom_margin = 2_000_000  # 2 см
    section.right_margin = 1_000_000   # 1 см

    # Настройка стиля по умолчанию
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = docx.shared.Pt(14)

    # Межстрочный интервал и отступ первой строки
    for paragraph in doc.paragraphs:
        if "{{faculty}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{faculty}}", faculty)

        paragraph.alignment = 3  # Выровнять по ширине

        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = docx.shared.Cm(1.25)
        paragraph_format.line_spacing = 1.5

        # Если заголовок, центрируем
        if paragraph.style.name.lower().startswith("heading"):
            paragraph.alignment = 1  # По центру
            paragraph_format.space_after = docx.shared.Pt(18)

    add_footer(doc, faculty)

    output_path = os.path.join(TEMP_DIR, generate_unique_filename())
    doc.save(output_path)
    logger.info(f"Сохраненный файл: {output_path}")
    return output_path

def cleanup_temp_dir():
    shutil.rmtree(TEMP_DIR, ignore_errors=True)

@app.post("/submit")
async def submit_file(
    background_tasks: BackgroundTasks,
    file: UploadFile,
    faculty: str = Form(...),
    session: AsyncSession = Depends(get_session)
):
    logger.info(f"Получен файл: {file.filename}")
    
    filename = file.filename.encode('utf-8').decode('utf-8')
    uploaded_file_path = os.path.join(TEMP_DIR, filename)
    
    with open(uploaded_file_path, "wb") as f:
        content = await file.read()
        f.write(content)

    template_path = uploaded_file_path
    logger.info(f"Заполнение шаблона: {template_path}")
    filled_template_path = fill_template(template_path, faculty)
    
    output_pdf_path = os.path.splitext(filled_template_path)[0] + ".pdf"
    logger.info(f"Конвертируем в PDF: {output_pdf_path}")
    convert(filled_template_path, output_pdf_path)

    # Обновляем записи в БД
    async with session.begin():
        # Обновление по факультету
        result = await session.execute(select(Faculty).filter_by(name=faculty))
        faculty_record = result.scalar()

        if faculty_record:
            faculty_record.submissions += 1
        else:
            faculty_record = Faculty(name=faculty, submissions=1)
            session.add(faculty_record)

        # ✅ Добавляем новую запись о подаче
        submission = Submission(faculty_name=faculty)
        session.add(submission)

    logger.info(f"Отправляем PDF: {output_pdf_path}")
    response = FileResponse(output_pdf_path, media_type='application/pdf', filename="coursework.pdf")

    # Добавляем задачу на удаление после отправки ответа
    background_tasks.add_task(cleanup_temp_dir)

    return response

# Эндпоинт для статистики
@app.get("/stats")
async def get_stats(session: AsyncSession = Depends(get_session)):
    result = await session.execute(select(Faculty))
    faculties = result.scalars().all()
    return [{"faculty": f.name, "submissions": f.submissions} for f in faculties]

@app.get("/stats_month")
async def get_monthly_stats(session: AsyncSession = Depends(get_session)):
    current_year = datetime.now().year

    result = await session.execute(
        select(
            extract('month', Submission.created_at).label("month"),
            func.count().label("count")
        )
        .where(extract('year', Submission.created_at) == current_year)
        .group_by("month")
        .order_by("month")
    )

    stats = result.all()

    # Преобразуем цифры месяцев в названия на русском
    month_names = [
        "Январь", "Февраль", "Март", "Апрель", "Май", "Июнь",
        "Июль", "Август", "Сентябрь", "Октябрь", "Ноябрь", "Декабрь"
    ]

    response = [
        {
            "date": f"{month_names[int(month)-1]} {current_year}",
            "count": count
        }
        for month, count in stats
    ]

    return response

if __name__ == "__main__":
    # Чтение порта и хоста из переменных окружения
    host = os.getenv("HOST", "0.0.0.0")  # 0.0.0.0 чтобы приложение было доступно извне
    port = int(os.getenv("PORT", 8000))  # По умолчанию порт 8000

    # Запуск сервера с заданными параметрами
    uvicorn.run("main:app", host=host, port=port, reload=True)